from __future__ import annotations

from pathlib import Path

from ai_dev_researcher.storage.docling_parser import parse_with_docling
from ai_dev_researcher.storage.normalized_docs import (
    _normalize_docx,
    _normalize_pdf,
    normalize_document,
)


def test_parse_with_docling_uses_formal_rapidocr_options_with_string_paths(
    monkeypatch,
):
    """Docling 的 RapidOCR 配置必须通过正式入口且不能携带 WindowsPath。"""
    import docling.document_converter as document_converter_module

    captured: dict[str, object] = {}

    class FakeConverter:
        def __init__(self, *args, **kwargs):
            captured["args"] = args
            captured["kwargs"] = kwargs

        def convert(self, path: str):
            class Page:
                text = "configured"

            class Document:
                pages = {1: Page()}

            class Result:
                document = Document()

            return Result()

    monkeypatch.setattr(document_converter_module, "DocumentConverter", FakeConverter)
    monkeypatch.setattr(
        "ai_dev_researcher.storage.torch_guard.ensure_torch_loaded", lambda: None
    )

    assert parse_with_docling(Path("sample.pdf")) == "[PAGE 1]\nconfigured"

    kwargs = captured["kwargs"]
    assert isinstance(kwargs, dict)
    format_options = kwargs["format_options"]
    pdf_option = next(iter(format_options.values()))
    pipeline_options = pdf_option.pipeline_options

    from docling.datamodel.pipeline_options import PdfPipelineOptions, RapidOcrOptions

    assert isinstance(pipeline_options, PdfPipelineOptions)
    assert isinstance(pipeline_options.ocr_options, RapidOcrOptions)

    rapidocr_options = pipeline_options.ocr_options
    for value in (
        rapidocr_options.det_model_path,
        rapidocr_options.cls_model_path,
        rapidocr_options.rec_model_path,
        rapidocr_options.rec_keys_path,
        rapidocr_options.font_path,
    ):
        assert value is None or isinstance(value, str)

    def assert_string_paths(value: object):
        if isinstance(value, dict):
            for item in value.values():
                assert_string_paths(item)
        elif isinstance(value, (list, tuple)):
            for item in value:
                assert_string_paths(item)
        else:
            assert not isinstance(value, Path)

    assert_string_paths(rapidocr_options.rapidocr_params)


def test_parse_with_docling_filters_only_omegaconf_merge_warning(monkeypatch):
    """只过滤已确认的 OmegaConf merge flag 告警，其他告警必须保留。"""
    import warnings

    import docling.document_converter as document_converter_module

    class FakeConverter:
        def __init__(self, *args, **kwargs):
            warnings.warn(
                "update() merge flag is is not specified, defaulting to False.\n"
                "For more details, see https://github.com/omry/omegaconf/issues/367",
                UserWarning,
            )
            warnings.warn("unrelated warning must remain visible", UserWarning)

        def convert(self, path: str):
            class Page:
                text = "configured"

            class Document:
                pages = {1: Page()}

            class Result:
                document = Document()

            return Result()

    monkeypatch.setattr(document_converter_module, "DocumentConverter", FakeConverter)
    monkeypatch.setattr(
        "ai_dev_researcher.storage.torch_guard.ensure_torch_loaded", lambda: None
    )

    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        assert parse_with_docling(Path("sample.pdf")) == "[PAGE 1]\nconfigured"

    messages = [str(item.message) for item in caught]
    assert messages == ["unrelated warning must remain visible"]


def test_parse_with_docling_returns_none_when_not_installed(monkeypatch, tmp_path: Path):
    """docling 未安装时返回 None（触发上层回退），不抛异常。"""
    import builtins

    real_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name.startswith("docling"):
            raise ImportError("docling not installed")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)
    pdf = tmp_path / "sample.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake")
    assert parse_with_docling(pdf) is None


def test_parse_with_docling_ignores_unsupported_ext(tmp_path: Path):
    txt = tmp_path / "note.txt"
    txt.write_text("hello", encoding="utf-8")
    assert parse_with_docling(txt) is None


def test_normalize_pdf_keeps_page_markers(tmp_path: Path):
    """pypdf 回退路径：输出应保留 [PAGE n] 标记。"""
    import pypdf
    from pypdf.generic import DecodedStreamObject, NameObject

    pdf_path = tmp_path / "two.pdf"
    writer = pypdf.PdfWriter()
    for text in ["first page", "second page"]:
        page = writer.add_blank_page(width=72, height=72)
        stream = DecodedStreamObject()
        stream.set_data(f"BT /F1 12 Tf 10 50 Td ({text}) Tj ET".encode("latin-1"))
        page[NameObject("/Contents")] = stream
    with pdf_path.open("wb") as fh:
        writer.write(fh)

    text = _normalize_pdf(pdf_path, max_pages=10)
    assert "[PAGE 1]" in text
    assert "[PAGE 2]" in text


def test_normalize_docx_keeps_paragraph_markers(tmp_path: Path):
    from docx import Document

    docx_path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("alpha")
    doc.add_paragraph("beta")
    doc.save(str(docx_path))

    text = _normalize_docx(docx_path)
    assert "[PARAGRAPH 1]" in text
    assert "alpha" in text
    assert "beta" in text


def test_normalize_document_docling_fallback_to_pypdf(tmp_path: Path, monkeypatch):
    """docling 不可用时，normalize_document 应回退到 pypdf 且不抛异常。"""
    import builtins

    real_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name.startswith("docling"):
            raise ImportError("docling not installed")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)

    import pypdf
    from pypdf.generic import DecodedStreamObject, NameObject

    pdf_path = tmp_path / "doc.pdf"
    writer = pypdf.PdfWriter()
    page = writer.add_blank_page(width=72, height=72)
    stream = DecodedStreamObject()
    stream.set_data(b"BT /F1 12 Tf 10 50 Td (fallback works) Tj ET")
    page[NameObject("/Contents")] = stream
    with pdf_path.open("wb") as fh:
        writer.write(fh)

    text = normalize_document(pdf_path, max_chars=10_000)
    assert "[PAGE 1]" in text
